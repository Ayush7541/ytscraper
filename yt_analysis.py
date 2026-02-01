# YouTube Channel Analysis Tool (Non-Shorts Only)
# Author: Ayush
# Requirements: google-api-python-client, isodate, pandas, tqdm

import os
import csv
import isodate
from datetime import datetime, timezone
from typing import List, Dict
import re
from googleapiclient.discovery import build
from tqdm import tqdm
import requests
from bs4 import BeautifulSoup
import time
import cloudscraper
# --- Excel support ---
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
import pandas as pd
# --- DOCX and Playwright ---
from playwright.sync_api import sync_playwright
from docx import Document

# ---------------- CONFIG ----------------

# ---- SET YOUR YOUTUBE API KEY HERE ----
YOUTUBE_API_KEY = "AIzaSyB2p8GXoG_AOm9TEDmqpoQiHoyO7IIe3eU"

API_SERVICE_NAME = "youtube"
API_VERSION = "v3"

# ---- CHANNELS TO ANALYZE (EDIT HERE) ----
# You can paste channel IDs, channel URLs, or @handles
CHANNELS_TO_ANALYZE = [
    "https://www.youtube.com/channel/UCYyel9YAgwvs07rz5kugU7g"
]

# ---- WEBSITES TO SCRAPE (EDIT HERE) ----
# Add any websites you want to scrape here
WEBSITES_TO_SCRAPE = [
    {
        "name": "",
        "url": ""
    },
    {
        "name": "",
        "url": ""
    }
]

# ---- TRANSCRIPT API CONFIG ----
TRANSCRIPT_API_KEY = "sk_nOhB6o3f9KdYsYe6orvSBhhMU7fCakH9rtfWptMY2DA"
TRANSCRIPT_API_BASE = "https://transcriptapi.com/api/v2/youtube/transcript"

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------- UTILS ----------------
def build_youtube():
    if not YOUTUBE_API_KEY or "PASTE_YOUR_YOUTUBE_API_KEY_HERE" in YOUTUBE_API_KEY:
        raise ValueError("Please add your YouTube API key directly in the code.")

    return build(API_SERVICE_NAME, API_VERSION, developerKey=YOUTUBE_API_KEY)


def parse_duration(duration_str: str) -> int:
    """Return duration in seconds"""
    return int(isodate.parse_duration(duration_str).total_seconds())


def is_non_short(video_item: Dict) -> bool:
    duration = parse_duration(video_item['contentDetails']['duration'])
    return duration >= 60


# ---------------- CORE DATA ----------------

def extract_channel_id(youtube, identifier: str) -> str:
    identifier = identifier.strip()

    # Case 1: Raw channel ID
    if identifier.startswith("UC"):
        return identifier

    # Case 2: Full channel URL
    if "youtube.com/channel/" in identifier:
        return identifier.split("youtube.com/channel/")[1].split("/")[0]

    # Case 3: Handle URL or raw handle
    if "youtube.com/@" in identifier:
        handle = identifier.split("youtube.com/@")[1]
    elif identifier.startswith("@"):
        handle = identifier[1:]
    else:
        handle = identifier

    # Resolve handle to channel ID via search
    res = youtube.search().list(
        part="snippet",
        q=handle,
        type="channel",
        maxResults=1
    ).execute()

    if not res.get("items"):
        raise ValueError(f"Could not resolve channel: {identifier}")

    return res["items"][0]["snippet"]["channelId"]


def get_channel_info(youtube, channel_id: str) -> Dict:
    res = youtube.channels().list(
        part="snippet,statistics",
        id=channel_id
    ).execute()

    item = res['items'][0]
    return {
        "channel_id": channel_id,
        "channel_name": item['snippet']['title'],
        "description": item['snippet']['description'],
        "subscribers": item['statistics'].get('subscriberCount'),
        "total_videos": item['statistics'].get('videoCount')
    }


# ---------------- VIDEOS ----------------
def get_uploads_playlist_id(youtube, channel_id: str) -> str:
    res = youtube.channels().list(
        part="contentDetails",
        id=channel_id
    ).execute()
    return res['items'][0]['contentDetails']['relatedPlaylists']['uploads']


def get_all_video_ids(youtube, playlist_id: str, max_results=200) -> List[str]:
    video_ids = []
    next_page = None

    while True:
        res = youtube.playlistItems().list(
            part="contentDetails",
            playlistId=playlist_id,
            maxResults=50,
            pageToken=next_page
        ).execute()

        for item in res['items']:
            video_ids.append(item['contentDetails']['videoId'])
            if len(video_ids) >= max_results:
                return video_ids

        next_page = res.get('nextPageToken')
        if not next_page:
            break

    return video_ids


def get_video_details(youtube, video_ids: List[str]) -> List[Dict]:
    all_items = []
    for i in range(0, len(video_ids), 50):
        chunk = video_ids[i:i+50]
        res = youtube.videos().list(
            part="snippet,statistics,contentDetails",
            id=",".join(chunk)
        ).execute()
        all_items.extend(res['items'])
    return all_items


# ---------------- ENGAGEMENT ----------------
def calculate_engagement(video: Dict) -> float:
    stats = video['statistics']
    views = int(stats.get('viewCount', 0))
    likes = int(stats.get('likeCount', 0))
    comments = int(stats.get('commentCount', 0))
    return round((likes + comments) / views, 6) if views > 0 else 0.0


# ---------------- COMMENTS ----------------
def get_comments(youtube, video_id: str) -> List[Dict]:
    comments = []
    next_page = None

    while True:
        res = youtube.commentThreads().list(
            part="snippet,replies",
            videoId=video_id,
            maxResults=100,
            pageToken=next_page,
            textFormat="plainText"
        ).execute()

        for item in res['items']:
            top = item['snippet']['topLevelComment']['snippet']
            comments.append({
                "video_id": video_id,
                "comment_id": item['id'],
                "parent_id": None,
                "author": top.get('authorDisplayName'),
                "text": top.get('textDisplay'),
                "likes": top.get('likeCount'),
                "published_at": top.get('publishedAt')
            })

            for reply in item.get('replies', {}).get('comments', []):
                rs = reply['snippet']
                comments.append({
                    "video_id": video_id,
                    "comment_id": reply['id'],
                    "parent_id": item['id'],
                    "author": rs.get('authorDisplayName'),
                    "text": rs.get('textDisplay'),
                    "likes": rs.get('likeCount'),
                    "published_at": rs.get('publishedAt')
                })

        next_page = res.get('nextPageToken')
        if not next_page:
            break

    return comments


# ---------------- WEB SCRAPER ----------------
def scrape_websites():
    results = []

    # ---- Tier 1: cloudscraper ----
    scraper = cloudscraper.create_scraper(
        browser={
            "browser": "chrome",
            "platform": "windows",
            "mobile": False
        }
    )

    for site in WEBSITES_TO_SCRAPE:
        url = site["url"]
        name = site["name"]

        # ---- URL validation guard ----
        if not url or not isinstance(url, str):
            continue

        url = url.strip()
        if not url.startswith(("http://", "https://")):
            continue

        try:
            res = scraper.get(url, timeout=30)
            res.raise_for_status()
            soup = BeautifulSoup(res.text, "html.parser")

            # remove scripts/styles
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()

            text = " ".join(soup.stripped_strings)

            # if meaningful text found, accept
            if text and len(text) > 500:
                results.append({
                    "source_name": name,
                    "url": url,
                    "method": "cloudscraper",
                    "text": text
                })
                continue

        except Exception:
            pass  # fallback to playwright

        # ---- Tier 2: Playwright fallback ----
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, timeout=60000)
            page.wait_for_load_state("networkidle")

            html = page.content()
            browser.close()

            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()

            text = " ".join(soup.stripped_strings)

            results.append({
                "source_name": name,
                "url": url,
                "method": "playwright",
                "text": text
            })

    return results


# ---------------- TRANSCRIPTS ----------------
def fetch_transcript(video_id: str) -> Dict:
    headers = {
        "Authorization": f"Bearer {TRANSCRIPT_API_KEY}"
    }

    params = {
        "video_url": video_id,
        "format": "text",
        "include_timestamp": False
    }

    max_retries = 3
    attempt = 0
    backoff_seconds = 2

    while attempt < max_retries:
        try:
            res = requests.get(
                TRANSCRIPT_API_BASE,
                headers=headers,
                params=params,
                timeout=30
            )

            if res.status_code == 200:
                data = res.json()
                return {
                    "video_id": data.get("video_id", video_id),
                    "language": data.get("language"),
                    "transcript": data.get("transcript", ""),
                    "status": "success",
                    "attempts": attempt + 1,
                    "fetched_at": datetime.now(timezone.utc).isoformat()
                }

            # Retryable errors
            if res.status_code in (408, 429, 503):
                attempt += 1
                if attempt < max_retries:
                    retry_after = res.headers.get("Retry-After")
                    wait_time = int(retry_after) if retry_after else backoff_seconds * attempt
                    time.sleep(wait_time)
                    continue

                return {
                    "video_id": video_id,
                    "language": None,
                    "transcript": "",
                    "status": f"retry_failed_{res.status_code}",
                    "attempts": attempt,
                    "fetched_at": datetime.now(timezone.utc).isoformat()
                }

            # Non-retryable errors
            return {
                "video_id": video_id,
                "language": None,
                "transcript": "",
                "status": f"error_{res.status_code}",
                "attempts": attempt + 1,
                "fetched_at": datetime.now(timezone.utc).isoformat()
            }

        except Exception as e:
            attempt += 1
            if attempt >= max_retries:
                return {
                    "video_id": video_id,
                    "language": None,
                    "transcript": "",
                    "status": f"exception: {e}",
                    "attempts": attempt,
                    "fetched_at": datetime.now(timezone.utc).isoformat()
                }
            time.sleep(backoff_seconds * attempt)

    return {
        "video_id": video_id,
        "language": None,
        "transcript": "",
        "status": "unknown_failure",
        "attempts": attempt,
        "fetched_at": datetime.now(timezone.utc).isoformat()
    }


# ---------------- CSV WRITERS ----------------
def write_csv(filename: str, rows: List[Dict]):
    if not rows:
        return
    path = os.path.join(OUTPUT_DIR, filename)
    with open(path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=rows[0].keys())
        writer.writeheader()
        writer.writerows(rows)


# ---------------- MASTER EXCEL WRITER ----------------
def write_master_excel(channel, latest_12, transcript_rows, comments_rows, website_rows):
    wb = Workbook()
    wb.remove(wb.active)

    # -------- CHANNEL SHEET --------
    channel_df = pd.DataFrame([channel])
    ws = wb.create_sheet("CHANNEL")
    for r in dataframe_to_rows(channel_df, index=False, header=True):
        ws.append(r)

    # -------- VIDEOS SHEET --------
    transcript_map = {t["video_id"]: t for t in transcript_rows}
    comments_map = {}
    for c in comments_rows:
        comments_map.setdefault(c["video_id"], []).append(c["text"])

    video_rows = []
    for v in latest_12:
        vid = v["id"]
        video_rows.append({
            "channel_id": channel["channel_id"],
            "channel_name": channel["channel_name"],
            "video_id": vid,
            "video_title": v["snippet"]["title"],
            "video_description": v["snippet"]["description"],
            "published_at": v["snippet"]["publishedAt"],
            "views": v["statistics"].get("viewCount"),
            "likes": v["statistics"].get("likeCount"),
            "comments_count": v["statistics"].get("commentCount"),
            "engagement": calculate_engagement(v),
            "transcript_status": transcript_map.get(vid, {}).get("status"),
            "transcript_text": transcript_map.get(vid, {}).get("transcript"),
            "all_comments_text": " || ".join(comments_map.get(vid, []))
        })

    videos_df = pd.DataFrame(video_rows)
    ws = wb.create_sheet("VIDEOS")
    for r in dataframe_to_rows(videos_df, index=False, header=True):
        ws.append(r)

    # -------- SCHEMA SHEET --------
    schema_rows = [
        {"field": "engagement", "definition": "(likes + comments_count) / views"},
        {"field": "transcript_text", "definition": "Transcript from TranscriptAPI v2. Plain text, no timestamps."},
        {"field": "all_comments_text", "definition": "All top-level comments and replies joined with ||"},
        {"field": "non_shorts", "definition": "Videos with duration >= 60 seconds"},
        {"field": "transcript_status", "definition": "success, retry_failed_*, or error_*"}
    ]
    schema_df = pd.DataFrame(schema_rows)
    ws = wb.create_sheet("SCHEMA")
    for r in dataframe_to_rows(schema_df, index=False, header=True):
        ws.append(r)

    wb.save(os.path.join(OUTPUT_DIR, "master.xlsx"))


# ------------- DOCX WRITER FOR WEBSITES -------------
def write_websites_docx(website_rows):
    doc = Document()
    doc.add_heading("Website Scraping Output", level=1)

    for w in website_rows:
        doc.add_heading(w["source_name"], level=2)
        doc.add_paragraph(f"URL: {w['url']}")
        doc.add_paragraph(f"Scraping method: {w['method']}")
        doc.add_paragraph("\n")
        doc.add_paragraph(w["text"])

    doc.save(os.path.join(OUTPUT_DIR, "websites.docx"))


def build_master_rows(channel, latest_12, transcripts, comments_rows, website_rows):
    # Index transcripts by video_id
    transcript_map = {t["video_id"]: t for t in transcripts}

    # Aggregate comments per video
    comments_map = {}
    for c in comments_rows:
        comments_map.setdefault(c["video_id"], []).append(c["text"])

    # Combine website text (same for all rows)
    website_text = " ".join([w.get("extracted_text", "") for w in website_rows])

    master_rows = []

    for v in latest_12:
        vid = v["id"]
        stats = v["statistics"]
        snippet = v["snippet"]

        master_rows.append({
            # Channel-level
            "channel_id": channel["channel_id"],
            "channel_name": channel["channel_name"],
            "channel_description": channel["description"],
            "subscribers": channel["subscribers"],
            "total_channel_videos": channel["total_videos"],

            # Video-level
            "video_id": vid,
            "video_title": snippet["title"],
            "video_description": snippet["description"],
            "published_at": snippet["publishedAt"],
            "views": stats.get("viewCount"),
            "likes": stats.get("likeCount"),
            "comments_count": stats.get("commentCount"),
            "engagement": calculate_engagement(v),

            # Transcript
            "transcript_status": transcript_map.get(vid, {}).get("status"),
            "transcript_text": transcript_map.get(vid, {}).get("transcript"),

            # Comments (aggregated)
            "all_comments_text": " || ".join(comments_map.get(vid, [])),

            # Websites
            "website_text": website_text,

            # ----- METADATA / LOGIC LABELS -----
            "engagement_definition": "(likes + comments_count) / views",
            "engagement_notes": "Calculated per video using public YouTube statistics. Views=0 returns engagement 0.0.",
            "comments_definition": "All top-level comments and replies aggregated into one field, separated by ' || '",
            "transcript_definition": "Transcript fetched via TranscriptAPI (v2). Plain text, no timestamps. Latest 4 non-short videos only.",
            "website_definition": "Text scraped from creator websites using cloudscraper to bypass WAF/ModSecurity.",
            "shorts_definition": "Video duration >= 60 seconds (YouTube ISO duration). Shorts excluded.",
            "master_row_definition": "Each row represents one non-short video from the latest 12 uploads of the channel."
        })

    return master_rows

# ---------------- MAIN PIPELINE ----------------
def analyze_channel(channel_identifier: str):
    youtube = build_youtube()

    channel_id = extract_channel_id(youtube, channel_identifier)

    # Channel
    channel = get_channel_info(youtube, channel_id)
    write_csv("channel.csv", [channel])

    # Videos
    playlist_id = get_uploads_playlist_id(youtube, channel_id)
    video_ids = get_all_video_ids(youtube, playlist_id)
    videos = get_video_details(youtube, video_ids)

    non_shorts = [v for v in videos if is_non_short(v)]

    # Latest 12 non-shorts
    latest_12 = sorted(non_shorts, key=lambda x: x['snippet']['publishedAt'], reverse=True)[:12]

    # Transcripts for latest 4 non-shorts
    transcript_rows = []
    for v in latest_12[:4]:
        transcript_rows.append(fetch_transcript(v['id']))

    write_csv("transcripts.csv", transcript_rows)

    video_rows = []
    for v in latest_12:
        video_rows.append({
            "video_id": v['id'],
            "title": v['snippet']['title'],
            "description": v['snippet']['description'],
            "views": v['statistics'].get('viewCount'),
            "likes": v['statistics'].get('likeCount'),
            "comments": v['statistics'].get('commentCount'),
            "engagement": calculate_engagement(v)
        })

    write_csv("latest_12_videos.csv", video_rows)

    # Engagement summary
    avg_engagement = round(sum([r['engagement'] for r in video_rows]) / len(video_rows), 6)
    write_csv("engagement_summary.csv", [{"avg_engagement_last_12": avg_engagement}])

    # Comments for last 30 non-shorts
    comments_rows = []
    last_30 = sorted(
        non_shorts,
        key=lambda x: x['snippet']['publishedAt'],
        reverse=True
    )[:30]

    for v in tqdm(last_30, desc="Fetching comments"):
        comments_rows.extend(get_comments(youtube, v['id']))

    write_csv("comments.csv", comments_rows)

    # Website scraping
    website_rows = scrape_websites()
    # write_csv("websites.csv", website_rows)   # REMOVED per instructions
    write_websites_docx(website_rows)

    # -------- MASTER CSV --------
    master_rows = build_master_rows(
        channel=channel,
        latest_12=latest_12,
        transcripts=transcript_rows,
        comments_rows=comments_rows,
        website_rows=website_rows
    )

    # REMOVE master.csv output per instructions
    # write_csv("master.csv", master_rows)

    # --- Write master Excel file ---
    write_master_excel(
        channel=channel,
        latest_12=latest_12,
        transcript_rows=transcript_rows,
        comments_rows=comments_rows,
        website_rows=website_rows
    )


if __name__ == "__main__":
    for channel in CHANNELS_TO_ANALYZE:
        analyze_channel(channel)
